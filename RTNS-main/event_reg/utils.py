from googleapiclient.discovery import build
from google.oauth2 import service_account
from googleapiclient.http import MediaIoBaseDownload
import io
from django.conf import settings
from docx import Document
import base64
def retrieve_file_content(file_id):
    credentials = service_account.Credentials.from_service_account_file(settings.GOOGLE_DRIVE_CREDENTIALS)
    drive_service = build('drive', 'v3', credentials=credentials)

    try:
        document_text = fetch_document_content(file_id)
        return document_text
    except Exception as e:
        print(e)

def fetch_document_content(file_id):
    credentials = service_account.Credentials.from_service_account_file(settings.GOOGLE_DRIVE_CREDENTIALS)
    drive_service = build('drive', 'v3', credentials=credentials)

    request = drive_service.files().get_media(fileId=file_id)
    file_content = io.BytesIO()
    downloader = MediaIoBaseDownload(file_content, request)
    done = False
    while done is False:
        status, done = downloader.next_chunk()
        
    document = Document(io.BytesIO(file_content.getvalue()))
    formatted_text = []
    for paragraph in document.paragraphs:
        for    run in paragraph.runs:
            text = run.text
            font = run.font
            bold = font.bold
            italic = font.italic
            underline = font.underline
            size = font.size
            formatted_text.append((text, bold, italic, underline, size))

    document_text = [paragraph.text for paragraph in document.paragraphs]

    return formatted_text








